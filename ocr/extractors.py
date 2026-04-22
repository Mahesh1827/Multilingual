"""
PDF & DOCX Text Extraction Module
Implements the smart text-layer-detection strategy:
  PDF → Text layer check (PyMuPDF)
   ├─ Text exists → Direct extraction
   └─ Scanned     → Convert to images for OCR

Enhancement: Structured table extraction from PDFs using PyMuPDF find_tables().
Tables are converted to Markdown + JSON for RAG-friendly storage.
"""

import pymupdf  # PyMuPDF
from pathlib import Path
from docx import Document as DocxDocument
from PIL import Image
import io
import json
import logging

logger = logging.getLogger(__name__)


def has_text_layer(page: pymupdf.Page, min_chars: int = 30) -> bool:
    """Check if a PDF page has an extractable text layer."""
    text = page.get_text("text").strip()
    return len(text) >= min_chars


# ──────────────────────────────────────────────
# Table Extraction
# ──────────────────────────────────────────────

def extract_tables_from_page(page: pymupdf.Page, page_num: int) -> list[dict]:
    """
    Extract tables from a PDF page using PyMuPDF's find_tables() API.
    Returns a list of dicts with both Markdown and JSON representations.
    """
    tables_data = []

    try:
        tables = page.find_tables()
        if not tables or len(tables.tables) == 0:
            return []

        for table_idx, table in enumerate(tables.tables):
            try:
                # Extract table data as list of lists
                rows = table.extract()
                if not rows or len(rows) < 2:
                    continue  # Skip tables with only headers or empty

                # Clean None values
                cleaned_rows = []
                for row in rows:
                    cleaned_rows.append([
                        (cell.strip() if cell else "") for cell in row
                    ])

                if not cleaned_rows:
                    continue

                headers = cleaned_rows[0]
                data_rows = cleaned_rows[1:]

                # Skip if all headers are empty
                if all(h == "" for h in headers):
                    if len(data_rows) > 0:
                        headers = [f"Column_{i+1}" for i in range(len(headers))]
                    else:
                        continue

                # Build Markdown table
                md_lines = []
                md_lines.append("| " + " | ".join(headers) + " |")
                md_lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
                for row in data_rows:
                    # Ensure row has same number of columns as headers
                    padded = row + [""] * (len(headers) - len(row))
                    md_lines.append("| " + " | ".join(padded[:len(headers)]) + " |")
                markdown_table = "\n".join(md_lines)

                # Build JSON representation
                json_rows = []
                for row in data_rows:
                    row_dict = {}
                    for i, header in enumerate(headers):
                        key = header if header else f"Column_{i+1}"
                        value = row[i] if i < len(row) else ""
                        row_dict[key] = value
                    json_rows.append(row_dict)

                # Create a natural language summary of the table
                summary = (
                    f"Table with {len(data_rows)} rows and {len(headers)} columns. "
                    f"Columns: {', '.join(h for h in headers if h)}."
                )

                # Combine into a chunk-ready text block
                table_text = (
                    f"{summary}\n\n"
                    f"{markdown_table}\n\n"
                    f"Structured data: {json.dumps(json_rows, ensure_ascii=False)}"
                )

                tables_data.append({
                    "page": page_num,
                    "table_index": table_idx,
                    "text": table_text,
                    "markdown": markdown_table,
                    "json_data": json_rows,
                    "headers": headers,
                    "num_rows": len(data_rows),
                    "method": "table_extraction",
                })

                logger.info(
                    f"Page {page_num}: Table {table_idx+1} extracted "
                    f"({len(data_rows)} rows × {len(headers)} cols)"
                )

            except Exception as e:
                logger.warning(f"Page {page_num}: Table {table_idx+1} extraction failed: {e}")
                continue

    except Exception as e:
        logger.debug(f"Page {page_num}: No tables or find_tables() failed: {e}")

    return tables_data


# ──────────────────────────────────────────────
# PDF Text Extraction
# ──────────────────────────────────────────────

def extract_text_from_pdf(pdf_path: Path) -> list[dict]:
    """
    Extract text and tables from a PDF using smart detection.
    Returns a list of dicts: {page, text, method, has_text_layer, tables}
    """
    results = []
    doc = pymupdf.open(str(pdf_path))

    for page_num in range(len(doc)):
        page = doc[page_num]
        page_data = {
            "page": page_num + 1,
            "text": "",
            "method": "",
            "has_text_layer": False,
            "tables": [],
        }

        # Extract tables first (works on all pages with text layers)
        if has_text_layer(page):
            # 🔥 Skip table detection on massive books to prevent pipeline freeze
            if len(doc) <= 60:
                page_tables = extract_tables_from_page(page, page_num + 1)
                page_data["tables"] = page_tables
            else:
                page_tables = []
                page_data["tables"] = []

            # Direct text extraction — fast, no OCR needed
            page_data["text"] = page.get_text("text").strip()
            page_data["method"] = "pymupdf_text"
            page_data["has_text_layer"] = True
            logger.info(
                f"Page {page_num + 1}: Text layer found → PyMuPDF extraction"
                f" ({len(page_tables)} tables detected)"
            )
        else:
            # No text layer — mark for OCR
            page_data["method"] = "needs_ocr"
            page_data["has_text_layer"] = False
            logger.info(f"Page {page_num + 1}: No text layer → will use OCR")

        results.append(page_data)

    doc.close()
    return results


def pdf_page_to_image(pdf_path: Path, page_num: int, dpi: int = 200) -> Image.Image:
    """Convert a single PDF page to a PIL Image for OCR."""
    import pymupdf
    from PIL import Image
    import io

    doc = pymupdf.open(str(pdf_path))
    page = doc[page_num]

    # safer rendering matrix
    zoom = dpi / 72
    mat = pymupdf.Matrix(zoom, zoom)

    pix = page.get_pixmap(matrix=mat, alpha=False)
    img_bytes = pix.tobytes("png")
    img = Image.open(io.BytesIO(img_bytes)).convert("RGB")

    doc.close()
    return img


def extract_text_from_docx(docx_path: Path) -> list[dict]:
    """Extract text and tables from a DOCX file."""
    doc = DocxDocument(str(docx_path))

    # Extract paragraphs
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())

    # Extract tables from DOCX
    tables_data = []
    for table_idx, table in enumerate(doc.tables):
        try:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)

            if len(rows) < 2:
                continue

            headers = rows[0]
            data_rows = rows[1:]

            # Build Markdown
            md_lines = []
            md_lines.append("| " + " | ".join(headers) + " |")
            md_lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
            for row in data_rows:
                padded = row + [""] * (len(headers) - len(row))
                md_lines.append("| " + " | ".join(padded[:len(headers)]) + " |")
            markdown_table = "\n".join(md_lines)

            # Build JSON
            json_rows = []
            for row in data_rows:
                row_dict = {}
                for i, header in enumerate(headers):
                    key = header if header else f"Column_{i+1}"
                    value = row[i] if i < len(row) else ""
                    row_dict[key] = value
                json_rows.append(row_dict)

            summary = (
                f"Table with {len(data_rows)} rows and {len(headers)} columns. "
                f"Columns: {', '.join(h for h in headers if h)}."
            )

            table_text = (
                f"{summary}\n\n"
                f"{markdown_table}\n\n"
                f"Structured data: {json.dumps(json_rows, ensure_ascii=False)}"
            )

            tables_data.append({
                "page": 1,
                "table_index": table_idx,
                "text": table_text,
                "markdown": markdown_table,
                "json_data": json_rows,
                "headers": headers,
                "num_rows": len(data_rows),
                "method": "table_extraction",
            })

        except Exception as e:
            logger.warning(f"DOCX table {table_idx} extraction failed: {e}")

    return [{
        "page": 1,
        "text": "\n".join(full_text),
        "method": "python_docx",
        "has_text_layer": True,
        "tables": tables_data,
    }]


def extract_from_file(file_path: Path) -> list[dict]:
    """Route extraction based on file type."""
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    else:
        logger.warning(f"Unsupported file type: {ext}")
        return []